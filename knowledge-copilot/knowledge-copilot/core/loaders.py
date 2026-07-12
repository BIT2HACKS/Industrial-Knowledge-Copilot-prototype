"""
loaders.py
Universal document loader for the Industrial Knowledge Copilot.
Supports: .txt, .pdf, .docx, .csv, and scanned images (.png/.jpg) via OCR.
This is the "Universal Document Ingestion" layer of the architecture -
it normalises every heterogeneous source format into plain text so the
same chunking/indexing pipeline can run over all of them.
"""
import os
import csv
import io

SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx", ".csv", ".png", ".jpg", ".jpeg"}


def load_txt(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def load_pdf(path):
    """Extract text page-by-page. Falls back to OCR per-page if a page has
    no extractable text (i.e. it's a scanned image, common for old
    inspection reports and regulatory filings)."""
    import pdfplumber

    text_parts = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text() or ""
            if not page_text.strip():
                page_text = _ocr_pdf_page(page)
            text_parts.append(f"[Page {i + 1}]\n{page_text}")
    return "\n\n".join(text_parts)


def _ocr_pdf_page(page):
    try:
        import pytesseract

        im = page.to_image(resolution=200).original
        return pytesseract.image_to_string(im)
    except Exception as e:
        return f"[OCR failed for this page: {e}]"


def load_docx(path):
    import docx

    doc = docx.Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def load_csv(path):
    rows = []
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        reader = csv.reader(f)
        for row in reader:
            rows.append(" | ".join(row))
    return "\n".join(rows)


def load_image(path):
    """OCR a scanned form or photographed document - e.g. a field
    technician's phone photo of a paper inspection checklist."""
    import pytesseract
    from PIL import Image

    return pytesseract.image_to_string(Image.open(path))


LOADERS = {
    ".txt": load_txt,
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".csv": load_csv,
    ".png": load_image,
    ".jpg": load_image,
    ".jpeg": load_image,
}


def load_document(path):
    ext = os.path.splitext(path)[1].lower()
    loader = LOADERS.get(ext)
    if not loader:
        raise ValueError(f"Unsupported file type: {ext}")
    return loader(path)


def discover_documents(folder):
    """Walk a folder and return paths of every supported document."""
    found = []
    for root, _, files in os.walk(folder):
        for name in files:
            ext = os.path.splitext(name)[1].lower()
            if ext in SUPPORTED_EXTENSIONS:
                found.append(os.path.join(root, name))
    return sorted(found)
